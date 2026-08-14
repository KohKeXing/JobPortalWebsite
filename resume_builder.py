#!/usr/bin/env python3
"""
Resume Portal - Python Companion Script
=======================================
This script is a fully-functional local companion to your Resume Portal web app.
It reads the exact JSON formats exported by the web interface, and allows you to:
1. Load, view, and save customized resumes.
2. Use Gemini AI (via the official Google GenAI Python SDK) to polish/improve descriptions.
3. Align and score your resume against any target Job Description.
4. Render and export gorgeous, fully responsive Single-Page HTML resumes matching 
   the Classic, Tech, Creative, and Sleek design systems of the portal.

Prerequisites:
  pip install google-genai

Usage:
  python resume_builder.py
"""

import os
import json
import sys
import uuid
import datetime
import html
import re
import zipfile
from io import BytesIO
from pathlib import Path

from pypdf import PdfReader
from pypdf.errors import PdfReadError
from typing import Dict, Any, List

from werkzeug.utils import secure_filename
from supabase_client import get_supabase_client
from file_encryption import encrypt_bytes, decrypt_bytes

# Try to import the modern official Google GenAI SDK
try:
    from google import genai
    from google.genai import types
    HAS_GEMINI = True
except ImportError:
    HAS_GEMINI = False

# Fallback values if Gemini SDK is not installed or API key is missing
FALLBACK_POLISH = (
    "• Spearheaded key strategic initiatives and delivered business-critical components.\n"
    "• Orchestrated scalable workflows and streamlined system integrations by 20%.\n"
    "• Collaborated with cross-functional squads to ensure pristine execution standards."
)

# ----------------------------------------------------
# DEFAULT / TEMPLATE DATA
# ----------------------------------------------------
DEFAULT_RESUME = {
    "personalInfo": {
        "name": "Alex Mercer",
        "email": "alex.mercer@techflow.io",
        "phone": "+1 (555) 342-9180",
        "location": "San Francisco, CA",
        "title": "Lead DevOps Platform Engineer",
        "summary": "Automation-focused Systems Architect and DevOps Lead with 5+ years of experience building resilient cloud systems. Expert in continuous orchestration, Kubernetes networking, and modern CI/CD patterns.",
        "website": "https://alexmercer.dev"
    },
    "experience": [
        {
            "id": "exp-1",
            "company": "CloudNet Labs",
            "role": "Lead DevOps Engineer",
            "startDate": "2024-01",
            "endDate": "Present",
            "description": "• Migrated legacy bare-metal systems to AWS cloud, slashing cloud bills by 35%.\n• Engineered secure continuous deployments via GitHub Actions and automated Terraform scripts."
        },
        {
            "id": "exp-2",
            "company": "CodeStream Inc",
            "role": "SRE Platform Specialist",
            "startDate": "2021-06",
            "endDate": "2023-12",
            "description": "• Managed 25+ production Kubernetes clusters with 99.99% operational uptime.\n• Implemented unified observability stack, reducing incident triage cycles by 50%."
        }
    ],
    "education": [
        {
            "id": "edu-1",
            "school": "UC Berkeley",
            "degree": "Bachelor of Science",
            "field": "Computer Science",
            "startDate": "2017-09",
            "endDate": "2021-05"
        }
    ],
    "skills": ["AWS", "Kubernetes", "Docker", "Terraform", "GitHub Actions", "Linux", "Python", "Go", "Datadog", "CI/CD"],
    "projects": [
        {
            "id": "proj-1",
            "name": "KubeDeploy Orchestrator",
            "description": "Lightweight CLI utility to boot multi-region cluster architectures in local dev sandboxes.",
            "technologies": ["Go", "Kubernetes", "Docker"]
        }
    ]
}


# ======================================================================
# WEB APP BACKEND STORAGE AND DOCUMENT GENERATION
# ======================================================================

RESUME_BUCKET = "resume-uploads"
MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_RESUME_PAGES = 10
MAX_COVER_LETTER_PAGES = 5
ALLOWED_RESUME_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_COVER_LETTER_EXTENSIONS = {".pdf", ".docx"}
CONTENT_TYPES = {
    ".pdf": "application/pdf",
    ".docx": (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    ),
}
ALLOWED_UPLOAD_MIME_TYPES = {
    ".pdf": {
        "application/pdf",
        "application/x-pdf",
        "application/octet-stream",
    },
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/zip",
        "application/x-zip-compressed",
        "application/octet-stream",
    },
}
MAX_DOCX_EXPANDED_BYTES = 50 * 1024 * 1024
MAX_DOCX_ARCHIVE_FILES = 1000
RESUME_COLUMNS = (
    "id,name,type,file_name,stored_file_name,file_format,storage_bucket,"
    "storage_path,layout,data,last_modified,owner_key"
)


def _utc_timestamp():
    return (
        datetime.datetime.now(datetime.timezone.utc)
        .isoformat()
        .replace("+00:00", "Z")
    )


def _api_resume(row):
    """Map Supabase columns to the field names already used by the UI."""
    record = {
        "id": row.get("id"),
        "name": row.get("name"),
        "type": row.get("type"),
        "lastModified": row.get("last_modified"),
        "ownerKey": row.get("owner_key"),
    }
    if row.get("file_name") or row.get("stored_file_name"):
        record.update(
            {
                "fileName": row.get("file_name"),
                "storedFileName": row.get("stored_file_name"),
                "fileFormat": row.get("file_format"),
            }
        )
    if row.get("type") == "builder":
        record.update(
            {
                "layout": row.get("layout") or "modern",
                "data": row.get("data") or {},
            }
        )
    return record


def _validate_pdf(content, max_pages):
    """Validate that an uploaded PDF is readable and within the page limit."""
    if not content.startswith(b"%PDF-"):
        raise ValueError(
            "The selected file is not a genuine PDF document."
        )

    try:
        reader = PdfReader(BytesIO(content))

        if reader.is_encrypted:
            raise ValueError(
                "Password-protected PDF files are not allowed."
            )

        page_count = len(reader.pages)
        print("Detected PDF pages:", page_count)

        if page_count == 0:
            raise ValueError(
                "The uploaded PDF does not contain any pages."
            )

        if page_count > max_pages:
            raise ValueError(
                f"The uploaded PDF has {page_count} pages. "
                f"The maximum allowed is {max_pages} pages."
            )

    except ValueError:
        raise
    except PdfReadError as exc:
        raise ValueError(
            "The uploaded PDF is corrupted or invalid."
        ) from exc
    except Exception as exc:
        raise ValueError(
            f"The uploaded PDF could not be read: {str(exc)}"
        ) from exc


def _validate_docx(content):
    """Validate the ZIP structure required by a genuine DOCX document."""
    if not content.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")):
        raise ValueError(
            "The selected file is not a genuine DOCX document."
        )

    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            entries = archive.infolist()
            names = {entry.filename for entry in entries}

            if len(entries) > MAX_DOCX_ARCHIVE_FILES:
                raise ValueError(
                    "The DOCX file contains too many internal files."
                )

            expanded_size = sum(entry.file_size for entry in entries)
            if expanded_size > MAX_DOCX_EXPANDED_BYTES:
                raise ValueError(
                    "The expanded DOCX content must not exceed 50 MB."
                )

            if any(entry.flag_bits & 0x1 for entry in entries):
                raise ValueError(
                    "Password-protected DOCX files are not allowed."
                )

            required_entries = {"[Content_Types].xml", "word/document.xml"}
            if not required_entries.issubset(names):
                raise ValueError(
                    "The selected file is not a valid Microsoft Word DOCX document."
                )

            broken_entry = archive.testzip()
            if broken_entry:
                raise ValueError(
                    "The DOCX file is corrupted or incomplete."
                )

    except ValueError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        raise ValueError(
            "The uploaded DOCX file is corrupted or invalid."
        ) from exc


def _read_upload(
    file_storage,
    allowed_extensions,
    max_pages=None,
):
    original_name = file_storage.filename or "uploaded_file"
    extension = Path(original_name).suffix.lower()

    if extension not in allowed_extensions:
        raise ValueError("Only PDF and DOCX files are supported.")

    if len(original_name) > 255:
        raise ValueError("The uploaded filename must not exceed 255 characters.")

    reported_mime = str(
        getattr(file_storage, "mimetype", "")
        or getattr(file_storage, "content_type", "")
        or ""
    ).split(";", 1)[0].strip().lower()

    if reported_mime.startswith("image/"):
        raise ValueError("Image files are not allowed. Upload a PDF or DOCX document.")
    if (
        reported_mime
        and reported_mime not in ALLOWED_UPLOAD_MIME_TYPES[extension]
    ):
        raise ValueError(
            "The file type does not match its PDF or DOCX extension."
        )

    safe_name = secure_filename(original_name)
    if not safe_name:
        safe_name = f"uploaded_file{extension}"

    file_storage.stream.seek(0)
    content = file_storage.stream.read()

    print("Upload filename:", safe_name)
    print("Upload extension:", extension)
    print("Upload size:", len(content), "bytes")

    if not content:
        raise ValueError("The uploaded file is empty.")

    if len(content) > MAX_UPLOAD_BYTES:
        raise ValueError("The uploaded file must not exceed 10 MB.")

    if extension == ".pdf":
        if max_pages is not None:
            _validate_pdf(content, max_pages)
    else:
        _validate_docx(content)

    # Reset the stream in case another part of the application needs it.
    file_storage.stream.seek(0)
    return safe_name, extension, content


def _clean_text(value):
    return str(value or "").strip()


def _description_lines(value):
    lines = []
    for line in _clean_text(value).splitlines():
        cleaned = re.sub(r"^[\s\u2022\-*]+", "", line).strip()
        if cleaned:
            lines.append(cleaned)
    return lines


def _output_format(value):
    output_format = _clean_text(value).lower().lstrip(".")
    if output_format not in {"pdf", "docx"}:
        raise ValueError("Choose either PDF or DOCX as the resume format.")
    return output_format


def _accent_hex(data, layout):
    defaults = {
        "modern": "#2563EB",
        "tech": "#0F766E",
        "elegant": "#7E22CE",
        "minimalist": "#334155",
    }
    accent = _clean_text((data or {}).get("accentColor"))
    if re.fullmatch(r"#[0-9a-fA-F]{6}", accent):
        return accent.upper()
    return defaults.get(layout, defaults["modern"])


def _generate_pdf_resume(name, layout, data):
    """Create a genuine, multi-page PDF resume in memory."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            KeepTogether,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:
        raise RuntimeError(
            "PDF generation requires reportlab. "
            "Run: python -m pip install reportlab"
        ) from exc

    data = data or {}
    info = data.get("personalInfo") or {}
    accent = colors.HexColor(_accent_hex(data, layout))
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=17 * mm,
        leftMargin=17 * mm,
        topMargin=15 * mm,
        bottomMargin=16 * mm,
        title=_clean_text(name) or "Resume",
        author=_clean_text(info.get("name")) or "JobPortal Candidate",
    )

    font_map = {
        "modern": ("Helvetica", "Helvetica-Bold"),
        "tech": ("Courier", "Courier-Bold"),
        "elegant": ("Times-Roman", "Times-Bold"),
        "minimalist": ("Times-Roman", "Times-Bold"),
    }
    body_font, bold_font = font_map.get(layout, font_map["modern"])
    centered_header = layout in {"elegant", "minimalist"}
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ResumeName",
        parent=styles["Title"],
        fontName=bold_font,
        fontSize=22,
        leading=26,
        textColor=accent,
        alignment=TA_CENTER if centered_header else TA_LEFT,
        spaceAfter=3,
    ))
    styles.add(ParagraphStyle(
        name="ResumeTitle",
        parent=styles["Normal"],
        fontName=body_font,
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#475569"),
        alignment=TA_CENTER if centered_header else TA_LEFT,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="ResumeContact",
        parent=styles["Normal"],
        fontName=body_font,
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#64748B"),
        alignment=TA_CENTER if centered_header else TA_LEFT,
        spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        name="ResumeSection",
        parent=styles["Heading2"],
        fontName=bold_font,
        fontSize=10,
        leading=13,
        textColor=accent,
        spaceBefore=8,
        spaceAfter=5,
        uppercase=True,
    ))
    styles.add(ParagraphStyle(
        name="ResumeBody",
        parent=styles["BodyText"],
        fontName=body_font,
        fontSize=9.2,
        leading=13,
        textColor=colors.HexColor("#334155"),
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="ResumeItemTitle",
        parent=styles["BodyText"],
        fontName=bold_font,
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor("#0F172A"),
    ))
    styles.add(ParagraphStyle(
        name="ResumeDate",
        parent=styles["BodyText"],
        fontName=body_font,
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#64748B"),
        alignment=2,
    ))
    styles.add(ParagraphStyle(
        name="ResumeBullet",
        parent=styles["BodyText"],
        fontName=body_font,
        fontSize=8.8,
        leading=12,
        leftIndent=10,
        firstLineIndent=-7,
        bulletIndent=0,
        textColor=colors.HexColor("#334155"),
        spaceAfter=2,
    ))

    story = []
    person_name = _clean_text(info.get("name")) or _clean_text(name)
    story.append(Paragraph(html.escape(person_name or "Candidate"), styles["ResumeName"]))
    if _clean_text(info.get("title")):
        story.append(Paragraph(
            html.escape(_clean_text(info.get("title"))),
            styles["ResumeTitle"],
        ))
    contact = [
        _clean_text(info.get("email")),
        _clean_text(info.get("phone")),
        _clean_text(info.get("location")),
        _clean_text(info.get("website")),
    ]
    contact = [html.escape(value) for value in contact if value]
    if contact:
        story.append(Paragraph(" &bull; ".join(contact), styles["ResumeContact"]))
    story.append(Table(
        [[""]],
        colWidths=[document.width],
        rowHeights=[1],
        style=TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), accent),
            ("LINEBELOW", (0, 0), (-1, -1), 0, accent),
        ]),
    ))
    story.append(Spacer(1, 5))

    def add_section(title):
        story.append(Paragraph(html.escape(title.upper()), styles["ResumeSection"]))

    summary = _clean_text(info.get("summary"))
    if summary:
        add_section("Professional Summary")
        story.append(Paragraph(
            html.escape(summary).replace("\n", "<br/>"),
            styles["ResumeBody"],
        ))

    experiences = data.get("experience") or []
    if experiences:
        add_section("Experience")
        for item in experiences:
            role = html.escape(_clean_text(item.get("role")) or "Role")
            company = html.escape(_clean_text(item.get("company")))
            dates = " - ".join(
                value for value in [
                    _clean_text(item.get("startDate")),
                    _clean_text(item.get("endDate")),
                ] if value
            )
            block = [
                Table(
                    [[
                        Paragraph(role, styles["ResumeItemTitle"]),
                        Paragraph(html.escape(dates), styles["ResumeDate"]),
                    ]],
                    colWidths=[document.width * 0.72, document.width * 0.28],
                    style=TableStyle([
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("TOPPADDING", (0, 0), (-1, -1), 1),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                    ]),
                ),
            ]
            if company:
                block.append(Paragraph(company, styles["ResumeBody"]))
            for line in _description_lines(item.get("description")):
                block.append(Paragraph(
                    html.escape(line),
                    styles["ResumeBullet"],
                    bulletText="\u2022",
                ))
            block.append(Spacer(1, 4))
            story.append(KeepTogether(block))

    education = data.get("education") or []
    if education:
        add_section("Education")
        for item in education:
            qualification = " in ".join(
                value for value in [
                    _clean_text(item.get("degree")),
                    _clean_text(item.get("field")),
                ] if value
            )
            school = _clean_text(item.get("school"))
            dates = " - ".join(
                value for value in [
                    _clean_text(item.get("startDate")),
                    _clean_text(item.get("endDate")),
                ] if value
            )
            story.append(KeepTogether([
                Table(
                    [[
                        Paragraph(
                            html.escape(qualification or "Qualification"),
                            styles["ResumeItemTitle"],
                        ),
                        Paragraph(html.escape(dates), styles["ResumeDate"]),
                    ]],
                    colWidths=[document.width * 0.72, document.width * 0.28],
                    style=TableStyle([
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("TOPPADDING", (0, 0), (-1, -1), 1),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                    ]),
                ),
                Paragraph(html.escape(school), styles["ResumeBody"]),
                Spacer(1, 3),
            ]))

    projects = data.get("projects") or []
    if projects:
        add_section("Key Projects")
        for item in projects:
            block = [
                Paragraph(
                    html.escape(_clean_text(item.get("name")) or "Project"),
                    styles["ResumeItemTitle"],
                ),
            ]
            if _clean_text(item.get("description")):
                block.append(Paragraph(
                    html.escape(_clean_text(item.get("description"))),
                    styles["ResumeBody"],
                ))
            technologies = [
                _clean_text(value)
                for value in (item.get("technologies") or [])
                if _clean_text(value)
            ]
            if technologies:
                block.append(Paragraph(
                    "<b>Technologies:</b> " + html.escape(", ".join(technologies)),
                    styles["ResumeBody"],
                ))
            block.append(Spacer(1, 3))
            story.append(KeepTogether(block))

    skills = [
        _clean_text(value)
        for value in (data.get("skills") or [])
        if _clean_text(value)
    ]
    if skills:
        add_section("Core Skills")
        story.append(Paragraph(
            " &bull; ".join(html.escape(value) for value in skills),
            styles["ResumeBody"],
        ))

    def draw_page(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#E2E8F0"))
        canvas.line(17 * mm, 11 * mm, A4[0] - 17 * mm, 11 * mm)
        canvas.setFont(body_font, 7.5)
        canvas.setFillColor(colors.HexColor("#94A3B8"))
        canvas.drawRightString(
            A4[0] - 17 * mm,
            7 * mm,
            f"Page {doc.page}",
        )
        canvas.restoreState()

    document.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    return buffer.getvalue()


def _generate_docx_resume(name, layout, data):
    """Create a genuine Microsoft Word DOCX resume in memory."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError(
            "DOCX generation requires python-docx. "
            "Run: python -m pip install python-docx"
        ) from exc

    data = data or {}
    info = data.get("personalInfo") or {}
    accent = _accent_hex(data, layout).lstrip("#")
    accent_color = RGBColor.from_string(accent)
    font_map = {
        "modern": "Aptos",
        "tech": "Consolas",
        "elegant": "Georgia",
        "minimalist": "Georgia",
    }
    font_name = font_map.get(layout, font_map["modern"])
    centered_header = layout in {"elegant", "minimalist"}

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def set_paragraph_spacing(paragraph, before=0, after=0, line=1.05):
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing = line

    def shade_paragraph(paragraph, fill):
        properties = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)

    def add_section_heading(title):
        paragraph = document.add_paragraph()
        set_paragraph_spacing(paragraph, before=7, after=3)
        run = paragraph.add_run(title.upper())
        run.bold = True
        run.font.name = font_name
        run.font.size = Pt(10)
        run.font.color.rgb = accent_color
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), accent)
        border.append(bottom)
        paragraph._p.get_or_add_pPr().append(border)

    header = document.add_paragraph()
    header.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if centered_header
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    set_paragraph_spacing(header, after=1)
    name_run = header.add_run(
        _clean_text(info.get("name")) or _clean_text(name) or "Candidate"
    )
    name_run.bold = True
    name_run.font.name = font_name
    name_run.font.size = Pt(21)
    name_run.font.color.rgb = accent_color

    if _clean_text(info.get("title")):
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = header.alignment
        set_paragraph_spacing(title_paragraph, after=2)
        title_run = title_paragraph.add_run(_clean_text(info.get("title")))
        title_run.font.name = font_name
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = RGBColor(71, 85, 105)

    contact_values = [
        _clean_text(info.get("email")),
        _clean_text(info.get("phone")),
        _clean_text(info.get("location")),
        _clean_text(info.get("website")),
    ]
    contact_values = [value for value in contact_values if value]
    if contact_values:
        contact = document.add_paragraph()
        contact.alignment = header.alignment
        set_paragraph_spacing(contact, after=5)
        contact_run = contact.add_run("  •  ".join(contact_values))
        contact_run.font.name = font_name
        contact_run.font.size = Pt(8.5)
        contact_run.font.color.rgb = RGBColor(100, 116, 139)

    accent_bar = document.add_paragraph()
    set_paragraph_spacing(accent_bar, after=3)
    shade_paragraph(accent_bar, accent)
    accent_bar.add_run(" ")

    summary = _clean_text(info.get("summary"))
    if summary:
        add_section_heading("Professional Summary")
        paragraph = document.add_paragraph(summary)
        set_paragraph_spacing(paragraph, after=3)

    experiences = data.get("experience") or []
    if experiences:
        add_section_heading("Experience")
        for item in experiences:
            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, before=2, after=0)
            role = paragraph.add_run(_clean_text(item.get("role")) or "Role")
            role.bold = True
            dates = " - ".join(
                value for value in [
                    _clean_text(item.get("startDate")),
                    _clean_text(item.get("endDate")),
                ] if value
            )
            if dates:
                date_run = paragraph.add_run(f"  |  {dates}")
                date_run.italic = True
                date_run.font.color.rgb = RGBColor(100, 116, 139)
            company_name = _clean_text(item.get("company"))
            if company_name:
                company = document.add_paragraph(company_name)
                set_paragraph_spacing(company, after=1)
                company.runs[0].font.color.rgb = accent_color
            for line in _description_lines(item.get("description")):
                bullet = document.add_paragraph(line, style="List Bullet")
                set_paragraph_spacing(bullet, after=0)

    education = data.get("education") or []
    if education:
        add_section_heading("Education")
        for item in education:
            qualification = " in ".join(
                value for value in [
                    _clean_text(item.get("degree")),
                    _clean_text(item.get("field")),
                ] if value
            )
            dates = " - ".join(
                value for value in [
                    _clean_text(item.get("startDate")),
                    _clean_text(item.get("endDate")),
                ] if value
            )
            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, before=2, after=0)
            paragraph.add_run(qualification or "Qualification").bold = True
            if dates:
                date_run = paragraph.add_run(f"  |  {dates}")
                date_run.italic = True
                date_run.font.color.rgb = RGBColor(100, 116, 139)
            school = document.add_paragraph(_clean_text(item.get("school")))
            set_paragraph_spacing(school, after=1)

    projects = data.get("projects") or []
    if projects:
        add_section_heading("Key Projects")
        for item in projects:
            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, before=2, after=0)
            paragraph.add_run(
                _clean_text(item.get("name")) or "Project"
            ).bold = True
            if _clean_text(item.get("description")):
                description = document.add_paragraph(
                    _clean_text(item.get("description"))
                )
                set_paragraph_spacing(description, after=0)
            technologies = [
                _clean_text(value)
                for value in (item.get("technologies") or [])
                if _clean_text(value)
            ]
            if technologies:
                technology = document.add_paragraph()
                set_paragraph_spacing(technology, after=1)
                technology.add_run("Technologies: ").bold = True
                technology.add_run(", ".join(technologies))

    skills = [
        _clean_text(value)
        for value in (data.get("skills") or [])
        if _clean_text(value)
    ]
    if skills:
        add_section_heading("Core Skills")
        skills_paragraph = document.add_paragraph("  •  ".join(skills))
        set_paragraph_spacing(skills_paragraph, after=1)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("JobPortal Resume")
    footer_run.font.name = font_name
    footer_run.font.size = Pt(7.5)
    footer_run.font.color.rgb = RGBColor(148, 163, 184)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _generate_builder_document(name, layout, data, output_format):
    output_format = _output_format(output_format)
    if output_format == "pdf":
        return _generate_pdf_resume(name, layout, data)
    return _generate_docx_resume(name, layout, data)


class ResumeStorage:
    """Supabase Database and private Storage CRUD for resumes."""

    def __init__(self, client=None):
        self._provided_client = client

    @property
    def client(self):
        return self._provided_client or get_supabase_client()

    def get_resumes(self):
        response = (
            self.client.table("resumes")
            .select(RESUME_COLUMNS)
            .order("last_modified", desc=True)
            .execute()
        )
        return [_api_resume(row) for row in (response.data or [])]

    def get_resume(self, resume_id):
        response = (
            self.client.table("resumes")
            .select(RESUME_COLUMNS)
            .eq("id", resume_id)
            .limit(1)
            .execute()
        )
        rows = response.data or []
        return _api_resume(rows[0]) if rows else None

    def add_uploaded_resume(self, file_storage, owner_key=None):
        safe_name, extension, content = _read_upload(
            file_storage,
            ALLOWED_RESUME_EXTENSIONS,
            MAX_RESUME_PAGES,
        )

        stored_name = f"{uuid.uuid4().hex}{extension}"
        storage_path = f"unassigned/{stored_name}"
        bucket = self.client.storage.from_(RESUME_BUCKET)

        print("Validation passed.")
        print("Uploading to bucket:", RESUME_BUCKET)
        print("Storage path:", storage_path)

        try:
            # 🔐 ENCRYPT BEFORE UPLOAD
            encrypted_content = encrypt_bytes(content)
            bucket.upload(
                path=storage_path,
                file=encrypted_content,
                file_options={
                    "content-type": CONTENT_TYPES[extension],
                    "upsert": "false",
                },
            )
            print("Supabase file upload successful.")
        except Exception as exc:
            print("SUPABASE STORAGE ERROR:", repr(exc))
            raise RuntimeError(
                f"Supabase Storage upload failed: {str(exc)}"
            ) from exc

        row = {
            "id": "res-" + uuid.uuid4().hex[:12],
            "name": Path(safe_name).stem,
            "type": "upload",
            "file_name": safe_name,
            "stored_file_name": stored_name,
            "file_format": extension.lstrip("."),
            "storage_bucket": RESUME_BUCKET,
            "storage_path": storage_path,
            "last_modified": _utc_timestamp(),
            "owner_key": owner_key,
        }

        try:
            response = self.client.table("resumes").insert(row).execute()
            print("Resume database record created successfully.")
        except Exception as exc:
            print("SUPABASE DATABASE ERROR:", repr(exc))
            try:
                bucket.remove([storage_path])
                print("Uploaded file removed after database failure.")
            except Exception as cleanup_exc:
                print("FILE CLEANUP ERROR:", repr(cleanup_exc))
            raise RuntimeError(
                f"Resume database insert failed: {str(exc)}"
            ) from exc

        rows = response.data or []
        if not rows:
            try:
                bucket.remove([storage_path])
            except Exception:
                pass
            raise RuntimeError(
                "The resume was uploaded, but no database record was returned."
            )

        return _api_resume(rows[0])

    def _store_generated_document(
        self,
        name,
        layout,
        data,
        output_format,
    ):
        output_format = _output_format(output_format)
        content = _generate_builder_document(
            name,
            layout,
            data,
            output_format,
        )
        if not content:
            raise RuntimeError("The generated resume file is empty.")
        if len(content) > MAX_UPLOAD_BYTES:
            raise ValueError("The generated resume must not exceed 10 MB.")

        safe_stem = secure_filename(
            Path(name or "resume").stem
        ) or "resume"
        extension = f".{output_format}"
        original_name = f"{safe_stem}{extension}"
        stored_name = f"{uuid.uuid4().hex}{extension}"
        storage_path = f"unassigned/{stored_name}"
        
        # 🔐 ENCRYPT BEFORE UPLOAD
        encrypted_content = encrypt_bytes(content)
        
        self.client.storage.from_(RESUME_BUCKET).upload(
            path=storage_path,
            file=encrypted_content,
            file_options={
                "content-type": CONTENT_TYPES[extension],
                "upsert": "false",
            },
        )
        return {
            "file_name": original_name,
            "stored_file_name": stored_name,
            "file_format": output_format,
            "storage_bucket": RESUME_BUCKET,
            "storage_path": storage_path,
        }

    def add_builder_resume(
        self,
        name,
        layout,
        data,
        output_format,
        owner_key=None,
    ):
        name = name or "Untitled Resume"
        layout = layout or "modern"
        data = data or {}
        stored_file = self._store_generated_document(
            name,
            layout,
            data,
            output_format,
        )
        row = {
            "id": "res-" + uuid.uuid4().hex[:12],
            "name": name,
            "type": "builder",
            "layout": layout,
            "data": data,
            **stored_file,
            "last_modified": _utc_timestamp(),
            "owner_key": owner_key,
        }
        try:
            response = self.client.table("resumes").insert(row).execute()
        except Exception:
            self.client.storage.from_(RESUME_BUCKET).remove(
                [stored_file["storage_path"]]
            )
            raise
        return _api_resume(response.data[0])

    def update_resume(self, resume_id, updates):
        existing_response = (
            self.client.table("resumes")
            .select(RESUME_COLUMNS)
            .eq("id", resume_id)
            .limit(1)
            .execute()
        )
        existing_rows = existing_response.data or []
        if not existing_rows:
            return None
        existing = existing_rows[0]

        field_map = {
            "name": "name",
            "layout": "layout",
            "data": "data",
        }
        changes = {
            database_field: updates[api_field]
            for api_field, database_field in field_map.items()
            if api_field in updates
        }
        changes["last_modified"] = _utc_timestamp()

        new_file = None
        builder_fields_changed = any(
            field in updates for field in ("name", "layout", "data", "outputFormat")
        )
        if existing.get("type") == "builder" and builder_fields_changed:
            output_format = (
                updates.get("outputFormat")
                or existing.get("file_format")
            )
            if not output_format:
                raise ValueError(
                    "Choose PDF or DOCX before saving this resume."
                )
            name = changes.get("name", existing.get("name"))
            layout = changes.get("layout", existing.get("layout") or "modern")
            data = changes.get("data", existing.get("data") or {})
            new_file = self._store_generated_document(
                name,
                layout,
                data,
                output_format,
            )
            changes.update(new_file)

        try:
            response = (
                self.client.table("resumes")
                .update(changes)
                .eq("id", resume_id)
                .execute()
            )
        except Exception:
            if new_file:
                self.client.storage.from_(RESUME_BUCKET).remove(
                    [new_file["storage_path"]]
                )
            raise

        old_storage_path = existing.get("storage_path")
        if (
            new_file
            and old_storage_path
            and old_storage_path != new_file["storage_path"]
        ):
            try:
                self.client.storage.from_(
                    existing.get("storage_bucket") or RESUME_BUCKET
                ).remove([old_storage_path])
            except Exception:
                pass

        rows = response.data or []
        if rows:
            return _api_resume(rows[0])
        return self.get_resume(resume_id)

    def delete_resume(self, resume_id):
        response = (
            self.client.table("resumes")
            .select(RESUME_COLUMNS)
            .eq("id", resume_id)
            .limit(1)
            .execute()
        )
        rows = response.data or []
        if not rows:
            return False

        row = rows[0]
        self.client.table("resumes").delete().eq("id", resume_id).execute()
        if row.get("storage_path"):
            self.client.storage.from_(
                row.get("storage_bucket") or RESUME_BUCKET
            ).remove([row["storage_path"]])
        return True

    def download_uploaded_resume(self, stored_filename):
        """Download and DECRYPT a resume file."""
        response = (
            self.client.table("resumes")
            .select(
                "id,file_name,file_format,storage_bucket,storage_path,"
                "stored_file_name,owner_key"
            )
            .eq("stored_file_name", stored_filename)
            .limit(1)
            .execute()
        )
        rows = response.data or []
        if not rows:
            return None

        row = rows[0]
        extension = "." + (row.get("file_format") or "pdf").lower()
        
        print(f"📁 Downloading: {row.get('file_name') or stored_filename}")
        
        # Download the encrypted file from Supabase
        encrypted_content = self.client.storage.from_(
            row.get("storage_bucket") or RESUME_BUCKET
        ).download(row["storage_path"])
        
        print(f"📦 Encrypted size: {len(encrypted_content)} bytes")
        
        # 🔑 DECRYPT THE FILE
        try:
            content = decrypt_bytes(encrypted_content)
            print(f"✅ Decrypted! Size: {len(content)} bytes")
        except Exception as e:
            print(f"❌ Decryption failed: {e}")
            # If decryption fails, keep the encrypted content
            content = encrypted_content
        
        return {
            "content": content,
            "file_name": row.get("file_name") or stored_filename,
            "content_type": CONTENT_TYPES.get(
                extension,
                "application/octet-stream",
            ),
            "resume_id": row.get("id"),
            "owner_key": row.get("owner_key"),
        }


def save_cover_letter_file(file_storage):
    """Upload a cover letter to the private resume bucket."""
    safe_name, extension, content = _read_upload(
        file_storage,
        ALLOWED_COVER_LETTER_EXTENSIONS,
        MAX_COVER_LETTER_PAGES,
    )
    stored_name = f"{uuid.uuid4().hex}{extension}"
    
    # 🔐 ENCRYPT BEFORE UPLOAD
    encrypted_content = encrypt_bytes(content)
    
    get_supabase_client().storage.from_(RESUME_BUCKET).upload(
        path=f"cover-letters/{stored_name}",
        file=encrypted_content,
        file_options={
            "content-type": CONTENT_TYPES[extension],
            "upsert": "false",
        },
    )
    return {
        "originalName": safe_name,
        "storedFileName": stored_name,
    }


def download_cover_letter_file(stored_filename):
    """Download and DECRYPT a cover letter file."""
    from file_encryption import decrypt_bytes
    
    extension = Path(stored_filename).suffix.lower()
    if extension not in ALLOWED_COVER_LETTER_EXTENSIONS:
        return None
    
    try:
        print(f"📁 Downloading cover letter: {stored_filename}")
        
        encrypted_content = get_supabase_client().storage.from_(RESUME_BUCKET).download(
            f"cover-letters/{stored_filename}"
        )
        
        print(f"📦 Encrypted size: {len(encrypted_content)} bytes")
        
        # 🔑 DECRYPT THE FILE
        try:
            content = decrypt_bytes(encrypted_content)
            print(f"✅ Cover letter decrypted!")
        except Exception as e:
            print(f"❌ Decryption failed: {e}")
            content = encrypted_content
            
    except Exception as e:
        print(f"❌ Download error: {e}")
        return None
    
    return {
        "content": content,
        "file_name": stored_filename,
        "content_type": CONTENT_TYPES.get(extension, "application/octet-stream"),
    }


def delete_cover_letter_file(stored_filename):
    """Delete a cover letter file from Supabase Storage."""
    if not stored_filename:
        return
    try:
        get_supabase_client().storage.from_(RESUME_BUCKET).remove(
            [f"cover-letters/{stored_filename}"]
        )
    except Exception:
        pass


# ======================================================================
# INTERACTIVE CLI COMPANION TOOL (kept for compatibility)
# ======================================================================

class ResumePortalPython:
    def __init__(self):
        self.resume_data = DEFAULT_RESUME.copy()
        self.client = None
        self.api_key_status = "Not Set"
        self.initialize_gemini()

    def initialize_gemini(self):
        api_key = os.environ.get("GEMINI_API_KEY")
        if not HAS_GEMINI:
            self.api_key_status = "Missing 'google-genai' library. Run 'pip install google-genai'"
            return

        if api_key:
            try:
                self.client = genai.Client(api_key=api_key)
                self.api_key_status = "Active"
            except Exception as e:
                self.api_key_status = f"Initialization Error: {str(e)}"
        else:
            self.api_key_status = "Missing GEMINI_API_KEY environment variable"

    def load_json(self, filepath: str) -> bool:
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = json.load(f)
            
            if "data" in content and isinstance(content["data"], dict):
                self.resume_data = content["data"]
            else:
                self.resume_data = content
            
            for field in ["personalInfo", "experience", "education", "skills", "projects"]:
                if field not in self.resume_data:
                    self.resume_data[field] = DEFAULT_RESUME[field] if field in DEFAULT_RESUME else []
            return True
        except Exception as e:
            print(f"\n[-] Error reading file: {str(e)}")
            return False

    def save_json(self, filepath: str) -> bool:
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(self.resume_data, f, indent=2, ensure_ascii=False)
            return True
        except Exception as e:
            print(f"\n[-] Error saving file: {str(e)}")
            return False

    def print_to_console(self):
        pi = self.resume_data.get("personalInfo", {})
        print("\n" + "="*60)
        print(f" {pi.get('name', 'N/A').upper()} - {pi.get('title', 'N/A')}")
        print(f" Email: {pi.get('email', 'N/A')} | Phone: {pi.get('phone', 'N/A')}")
        print(f" Location: {pi.get('location', 'N/A')} | Website: {pi.get('website', 'N/A')}")
        print("="*60)
        
        summary = pi.get("summary", "")
        if summary:
            print("\n[EXECUTIVE SUMMARY]")
            print(summary)
            
        print("\n[PROFESSIONAL EXPERIENCE]")
        for exp in self.resume_data.get("experience", []):
            print(f"• {exp.get('role')} @ {exp.get('company')} ({exp.get('startDate')} - {exp.get('endDate')})")
            desc = exp.get("description", "")
            for line in desc.strip().split("\n"):
                print(f"  {line}")
                
        print("\n[EDUCATION]")
        for edu in self.resume_data.get("education", []):
            print(f"• {edu.get('degree')} in {edu.get('field')} - {edu.get('school')} ({edu.get('startDate')} - {edu.get('endDate')})")
            
        print("\n[SKILLS]")
        print(", ".join(self.resume_data.get("skills", [])))
        
        print("\n[KEY PROJECTS]")
        for proj in self.resume_data.get("projects", []):
            techs = ", ".join(proj.get("technologies", []))
            print(f"• {proj.get('name')} [{techs}]")
            print(f"  {proj.get('description')}")
        print("="*60 + "\n")


def main():
    portal = ResumePortalPython()
    
    while True:
        print("\n==================================================")
        print("    RESUME PORTAL - PYTHON COMPANION SERVICE      ")
        print("==================================================")
        print(f" Gemini API Status: {portal.api_key_status}")
        print("--------------------------------------------------")
        print(" 1. View Current Resume Records")
        print(" 2. Load Resume JSON File (from Web App Export)")
        print(" 3. Save Resume JSON File")
        print(" 4. Exit Console")
        print("==================================================")
        
        choice = input("Select an option (1-4): ").strip()
        
        if choice == "1":
            portal.print_to_console()
        elif choice == "2":
            path_in = input("\nEnter path to downloaded resume JSON file: ").strip()
            if not path_in:
                path_in = "my_resume_data.json"
            if portal.load_json(path_in):
                print(f"\n[+] Successfully loaded data from '{path_in}'!")
                portal.print_to_console()
        elif choice == "3":
            path_out = input("\nEnter destination JSON filename (default: exported_resume.json): ").strip()
            if not path_out:
                path_out = "exported_resume.json"
            if portal.save_json(path_out):
                print(f"\n[+] Successfully saved resume records to '{path_out}'!")
        elif choice == "4":
            print("\nThank you for using the Resume Portal Companion!")
            sys.exit(0)
        else:
            print("\n[-] Invalid option. Please select 1 through 4.")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\nProcess interrupted. Goodbye!")
        sys.exit(0)
